"""处理批量导入和预览的函数"""

import streamlit as st
from typing import Dict, Any, List, Optional
import pandas as pd
from docx import Document
import PyPDF2
from io import BytesIO, StringIO
import re
# from test_batch import BatchProcessor  # 暂时注释掉，模块不存在
from ai_requirement_processor import AIRequirementProcessor, estimate_requirement_complexity

# 常量
MIN_PARAGRAPH_LENGTH = 10

def fetch_feishu_document(doc_id_or_url: str) -> Optional[str]:
    """获取飞书文档内容"""
    try:
        # 导入test.py中的fetch_url_content函数
        from test import fetch_url_content
        
        # 如果是文档ID而不是完整URL，构造URL
        if not doc_id_or_url.startswith('http'):
            doc_id_or_url = f"https://mi.feishu.cn/docx/{doc_id_or_url}"
        
        # 使用test.py中的fetch_url_content函数
        content = fetch_url_content(doc_id_or_url)
        
        # 检查是否返回错误 - 只有当内容明确以错误标识开始时才返回None
        if content.startswith("【飞书API错误】") or "网页抓取但需要登录" in content:
            return None
            
        return content
        
    except Exception as e:
        import logging
        logging.error(f"获取飞书文档失败: {e}")
        return None

def handle_batch_input() -> None:
    """处理批量导入需求的输入部分"""
    try:
        st.markdown("### 需求输入")
        
        # AI需求处理配置
        st.markdown("#### AI需求智能处理")
        col1, col2 = st.columns(2)
        with col1:
            enable_ai_analysis = st.checkbox("启用AI需求分析", value=True, 
                                           help="使用AI自动识别需求类型、优先级和复杂度")
        with col2:
            enable_ai_decomposition = st.checkbox("启用AI需求分解", value=True,
                                                help="自动将复杂需求分解为可测试的子需求")
        
        # 清空按钮
        if st.button("🗑️ 清空所有需求"):
            st.session_state.collected_requirements = []
            st.session_state.source_counts = []
            st.session_state.enable_ai_analysis = enable_ai_analysis
            st.session_state.enable_ai_decomposition = enable_ai_decomposition
            st.success("已清空所有需求")
        
        # 1. 飞书文档输入
        feishu_doc = st.text_input(
            "飞书文档链接或ID", 
            placeholder="输入飞书文档链接或ID"
        )
        if feishu_doc:
            with st.spinner("正在读取飞书文档..."):
                doc_content = fetch_feishu_document(feishu_doc)
                if doc_content:
                    parts = re.split(r"\n\s*\n+", doc_content.strip())
                    feishu_reqs = [p for p in parts 
                                if len(p.strip()) > MIN_PARAGRAPH_LENGTH]
                    if feishu_reqs:
                        add_requirements_batch(feishu_reqs, "飞书文档", 
                                             enable_ai_analysis, enable_ai_decomposition)
                        st.success(f"已导入 {len(feishu_reqs)} 条需求")
        
        # 2. 文件上传
        uploaded_files = st.file_uploader(
            "上传需求文件",
            type=["xlsx", "docx", "pdf", "txt", "csv"],
            accept_multiple_files=True
        )
        
        if uploaded_files:
            for file in uploaded_files:
                with st.spinner(f"正在处理 {file.name}..."):
                    process_uploaded_file(file, enable_ai_analysis, enable_ai_decomposition)
        
        # 3. 手动输入
        manual_reqs = st.text_area(
            "直接输入需求（每行一条）",
            placeholder="需求1\n需求2\n需求3...",
            height=150
        )
        
        if st.button("添加手工输入"):
            if manual_reqs:
                lines = [l.strip() for l in manual_reqs.splitlines() 
                        if len(l.strip()) > MIN_PARAGRAPH_LENGTH]
                if lines:
                    add_requirements_batch(lines, "手工输入", 
                                         enable_ai_analysis, enable_ai_decomposition)
                    st.success(f"已添加 {len(lines)} 条需求")
            else:
                st.warning("请输入需求内容")
                
    except Exception as e:
        st.error(f"需求输入处理错误: {str(e)}")
        if st.session_state.get("debug_mode"):
            st.exception(e)

def handle_batch_preview_and_generate(
    base_url: str,
    model: str,
    temperature: float,
    headers: List[str],
    pos_n: int,
    neg_n: int,
    edge_n: int,
    auto_mode: bool,
    dyn_params: Dict[str, Any],
    api_key: str
) -> None:
    """处理批量需求的预览和生成部分"""
    try:
        st.markdown("### 需求预览与生成")
        
        # 获取已收集的需求
        requirements = st.session_state.get("collected_requirements", [])
        
        if not requirements:
            st.warning("请先添加需求")
            return
        
        # 显示统计信息
        source_counts = st.session_state.get("source_counts", [])
        if source_counts:
            st.info("数据来源: " + " | ".join(source_counts))
        st.info(f"总计: {len(requirements)} 条需求")
        
        # AI需求处理选项
        enable_ai_analysis = st.session_state.get("enable_ai_analysis", True)
        enable_ai_decomposition = st.session_state.get("enable_ai_decomposition", True)
        
        if enable_ai_analysis and api_key:
            st.markdown("#### AI需求智能处理")
            if st.button("🔍 执行AI需求分析", type="secondary"):
                with st.spinner("正在执行AI需求分析..."):
                    try:
                        # 创建AI处理器
                        import uuid
                        ai_processor = AIRequirementProcessor(
                            client=OpenAI(
                                api_key=api_key,
                                base_url=str(base_url or "").rstrip("/") + "/v1",
                                default_headers={
                                    "X-Model-Provider-Id": "xiaomi",
                                    "X-Model-Request-Id": str(uuid.uuid4()),
                                },
                            ),
                            model=model,
                            temperature=temperature,
                        )
                        
                        # 获取需求文本列表
                        req_texts = [r["需求描述"] for r in requirements]
                        
                        # 执行AI分析
                        processed_reqs = ai_processor.process_batch_requirements(req_texts)
                        
                        # 更新需求信息
                        for i, processed_req in enumerate(processed_reqs):
                            if i < len(requirements):
                                requirements[i].update({
                                    "类型": processed_req["type"],
                                    "优先级": processed_req["priority"],
                                    "复杂度": processed_req["complexity"],
                                    "是否分解": "是" if processed_req["is_decomposed"] else "否"
                                })
                        
                        st.success(f"AI分析完成！共分析 {len(processed_reqs)} 条需求")
                        
                    except Exception as e:
                        st.error(f"AI需求分析失败: {str(e)}")
        
        # 预览表格
        preview_df = pd.DataFrame(requirements)
        st.dataframe(preview_df, use_container_width=True)
        
        st.divider()
        st.markdown("### 生成设置")
        
        parallel = st.number_input("并行处理数", 1, 8, 4)
        progress_ph = st.empty()
        result_ph = st.empty()
        
        if st.button("开始批量生成", type="primary"):
            try:
                progress_bar = progress_ph.progress(0)
                result_ph.info("正在生成...")
                
                # 创建处理器
                processor = BatchProcessor(
                    model=model,
                    base_url=base_url,
                    headers=headers,
                    pos_n=pos_n,
                    neg_n=neg_n,
                    edge_n=edge_n,
                    temperature=temperature,
                    max_workers=parallel,
                    background_knowledge=st.session_state.get('background_knowledge'),
                    dynamic_mode=auto_mode,
                    dynamic_params=dyn_params
                )
                
                # 准备需求列表
                req_list = get_requirements_for_batch(requirements)
                
                # 如果启用了AI分解，使用分解后的子需求
                if enable_ai_decomposition and api_key:
                    try:
                        import uuid
                        ai_processor = AIRequirementProcessor(
                            client=OpenAI(
                                api_key=api_key,
                                base_url=str(base_url or "").rstrip("/") + "/v1",
                                default_headers={
                                    "X-Model-Provider-Id": "xiaomi",
                                    "X-Model-Request-Id": str(uuid.uuid4()),
                                },
                            ),
                            model=model,
                            temperature=temperature,
                        )
                        
                        # 对复杂需求进行分解
                        decomposed_reqs = []
                        for req_text in req_list:
                            complexity = estimate_requirement_complexity(req_text)
                            if complexity == "高":
                                sub_reqs = ai_processor.decompose_requirement(req_text)
                                for sub_req in sub_reqs:
                                    decomposed_reqs.append(sub_req["sub_requirement"])
                            else:
                                decomposed_reqs.append(req_text)
                        
                        req_list = decomposed_reqs
                        st.info(f"AI分解后共 {len(req_list)} 条可测试需求")
                        
                    except Exception as e:
                        st.warning(f"AI需求分解失败，使用原始需求: {str(e)}")
                
                # 执行生成
                df_result = processor.process_batch(req_list)
                progress_bar.progress(100)
                
                if df_result is not None and not df_result.empty:
                    result_ph.success(f"已生成 {len(df_result)} 条测试用例")
                    st.dataframe(df_result, use_container_width=True)
                    
                    # 准备下载
                    excel_data = BytesIO()
                    with pd.ExcelWriter(excel_data, engine='openpyxl') as writer:
                        df_result.to_excel(writer, index=False)
                    excel_data.seek(0)
                    
                    # 下载按钮
                    st.download_button(
                        "📥 下载 Excel",
                        data=excel_data,
                        file_name="测试用例_批量.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    )
                    
                    st.download_button(
                        "📥 下载 CSV",
                        data=df_result.to_csv(index=False).encode('utf-8-sig'),
                        file_name="测试用例_批量.csv",
                        mime="text/csv"
                    )
                    
                    # 错误信息
                    errors = processor.get_errors()
                    if errors:
                        with st.expander(f"处理过程中的错误 ({len(errors)})"):
                            for req_id, error in errors:
                                st.error(f"{req_id}: {error}")
                else:
                    result_ph.error("生成失败，未获得有效结果")
                    
            except Exception as e:
                progress_ph.empty()
                result_ph.error(f"批量生成失败: {str(e)}")
                if st.session_state.get("debug_mode"):
                    st.exception(e)
                    
    except Exception as e:
        st.error(f"预览和生成错误: {str(e)}")
        if st.session_state.get("debug_mode"):
            st.exception(e)

def process_uploaded_file(file, enable_ai_analysis: bool = True, enable_ai_decomposition: bool = True) -> None:
    """处理上传的文件，使用AI智能识别需求"""
    try:
        name = file.name.lower()
        
        if name.endswith('.xlsx'):
            df = pd.read_excel(file)
            # 自动检测需求列
            req_columns = [col for col in df.columns if any(keyword in col for keyword in ['需求', '要求', '功能', '描述'])]
            if not req_columns:
                req_columns = [df.columns[0]]  # 默认使用第一列
            
            col = st.selectbox(f"选择需求列 ({file.name})", req_columns)
            df_reqs = df[col].dropna().astype(str).str.strip().tolist()
            
            if df_reqs:
                add_requirements_batch(df_reqs, f"Excel-{file.name}", 
                                     enable_ai_analysis, enable_ai_decomposition)
                st.success(f"已导入 {len(df_reqs)} 条需求")
        
        elif name.endswith('.docx'):
            doc = Document(file)
            # 提取所有段落文本
            full_text = "\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])
            
            if enable_ai_analysis:
                # 使用AI智能识别需求段落
                st.info("使用AI智能识别Word文档中的需求...")
                ai_identified_reqs = identify_requirements_with_ai(full_text, file.name)
                if ai_identified_reqs:
                    add_requirements_batch(ai_identified_reqs, f"Word-AI识别-{file.name}",
                                         enable_ai_analysis, enable_ai_decomposition)
                    st.success(f"AI识别出 {len(ai_identified_reqs)} 条需求")
                else:
                    # AI识别失败，使用传统方法
                    word_reqs = [p.text.strip() for p in doc.paragraphs 
                                if len(p.text.strip()) > MIN_PARAGRAPH_LENGTH]
                    if word_reqs:
                        add_requirements_batch(word_reqs, f"Word-传统-{file.name}",
                                             enable_ai_analysis, enable_ai_decomposition)
                        st.success(f"传统方法导入 {len(word_reqs)} 条需求")
            else:
                # 不使用AI，使用传统方法
                word_reqs = [p.text.strip() for p in doc.paragraphs 
                            if len(p.text.strip()) > MIN_PARAGRAPH_LENGTH]
                if word_reqs:
                    add_requirements_batch(word_reqs, f"Word-传统-{file.name}",
                                         enable_ai_analysis, enable_ai_decomposition)
                    st.success(f"已导入 {len(word_reqs)} 条需求")
        
        elif name.endswith('.pdf'):
            pdf_reader = PyPDF2.PdfReader(BytesIO(file.getvalue()))
            # 提取所有文本
            full_text = ""
            for page in pdf_reader.pages:
                full_text += page.extract_text() + "\n"
            
            if enable_ai_analysis:
                # 使用AI智能识别PDF中的需求
                st.info("使用AI智能识别PDF文档中的需求...")
                ai_identified_reqs = identify_requirements_with_ai(full_text, file.name)
                if ai_identified_reqs:
                    add_requirements_batch(ai_identified_reqs, f"PDF-AI识别-{file.name}",
                                         enable_ai_analysis, enable_ai_decomposition)
                    st.success(f"AI识别出 {len(ai_identified_reqs)} 条需求")
                else:
                    # AI识别失败，使用传统方法
                    parts = re.split(r"\n\s*\n+", full_text.strip())
                    pdf_reqs = [p for p in parts if len(p.strip()) > MIN_PARAGRAPH_LENGTH]
                    if pdf_reqs:
                        add_requirements_batch(pdf_reqs, f"PDF-传统-{file.name}",
                                             enable_ai_analysis, enable_ai_decomposition)
                        st.success(f"传统方法导入 {len(pdf_reqs)} 条需求")
            else:
                # 不使用AI，使用传统方法
                parts = re.split(r"\n\s*\n+", full_text.strip())
                pdf_reqs = [p for p in parts if len(p.strip()) > MIN_PARAGRAPH_LENGTH]
                if pdf_reqs:
                    add_requirements_batch(pdf_reqs, f"PDF-传统-{file.name}",
                                         enable_ai_analysis, enable_ai_decomposition)
                    st.success(f"已导入 {len(pdf_reqs)} 条需求")
        
        elif name.endswith(('.txt', '.csv')):
            stringio = StringIO(file.getvalue().decode("utf-8"))
            full_text = stringio.read()
            
            if enable_ai_analysis:
                # 使用AI智能识别文本中的需求
                st.info("使用AI智能识别文本文件中的需求...")
                ai_identified_reqs = identify_requirements_with_ai(full_text, file.name)
                if ai_identified_reqs:
                    add_requirements_batch(ai_identified_reqs, f"Text-AI识别-{file.name}",
                                         enable_ai_analysis, enable_ai_decomposition)
                    st.success(f"AI识别出 {len(ai_identified_reqs)} 条需求")
                else:
                    # AI识别失败，使用传统方法
                    lines = [l.strip() for l in full_text.splitlines() 
                            if len(l.strip()) > MIN_PARAGRAPH_LENGTH]
                    if lines:
                        add_requirements_batch(lines, f"Text-传统-{file.name}",
                                             enable_ai_analysis, enable_ai_decomposition)
                        st.success(f"传统方法导入 {len(lines)} 条需求")
            else:
                # 不使用AI，使用传统方法
                lines = [l.strip() for l in full_text.splitlines() 
                        if len(l.strip()) > MIN_PARAGRAPH_LENGTH]
                if lines:
                    add_requirements_batch(lines, f"Text-传统-{file.name}",
                                         enable_ai_analysis, enable_ai_decomposition)
                    st.success(f"已导入 {len(lines)} 条需求")
                
    except Exception as e:
        st.error(f"处理文件 {file.name} 失败: {str(e)}")
        if st.session_state.get("debug_mode"):
            st.exception(e)

def add_requirements_batch(requirements: List[str], source: str, 
                          enable_ai_analysis: bool = True, enable_ai_decomposition: bool = True) -> None:
    """添加一批需求到会话状态"""
    if not hasattr(st.session_state, "collected_requirements"):
        st.session_state.collected_requirements = []
    if not hasattr(st.session_state, "source_counts"):
        st.session_state.source_counts = []
    
    # 保存AI处理配置
    st.session_state.enable_ai_analysis = enable_ai_analysis
    st.session_state.enable_ai_decomposition = enable_ai_decomposition
        
    for req in requirements:
        req_text = req.strip()
        if not req_text:
            continue
            
        # 基础分析（即使AI分析未启用）
        complexity = estimate_requirement_complexity(req_text)
        
        st.session_state.collected_requirements.append({
            "需求编号": "",
            "需求描述": req_text,
            "来源": source,
            "复杂度": complexity,
            "类型": "待分析",
            "优先级": "待分析",
            "是否分解": "否"
        })
    st.session_state.source_counts.append(f"{source}:{len(requirements)}")

def get_requirements_for_batch(requirements: List[Dict[str, str]]) -> List[str]:
    """将需求列表转换为批处理格式"""
    return [r["需求描述"] for r in requirements if r["需求描述"].strip()]

def identify_requirements_with_ai(full_text: str, filename: str) -> List[str]:
    """使用AI智能识别文档中的需求"""
    try:
        # 检查是否有API配置
        if not st.session_state.get('api_key') or not st.session_state.get('base_url'):
            st.warning("未配置API Key，无法使用AI需求识别")
            return []
        
        # 创建AI处理器
        from openai import OpenAI
        import uuid
        client = OpenAI(
            api_key=st.session_state.get('api_key'),
            base_url=str(st.session_state.get('base_url') or "").rstrip("/") + "/v1",
            default_headers={
                "X-Model-Provider-Id": "xiaomi",
                "X-Model-Request-Id": str(uuid.uuid4()),
            },
        )
        
        # 构建提示词
        prompt = f"""请从以下文档内容中识别出所有的软件需求。文档内容：
        
{full_text[:4000]}  # 限制文本长度避免token超限

请按照以下要求识别需求：
1. 识别独立的功能需求、性能需求、安全需求等
2. 每个需求应该是完整、可测试的独立单元
3. 忽略文档的格式标记、标题、页眉页脚等非需求内容
4. 将识别出的需求按JSON数组格式返回

返回格式：
{{
    "requirements": [
        "需求1描述",
        "需求2描述",
        ...
    ]
}}

请只返回JSON格式，不要有其他内容。"""
        
        response = client.chat.completions.create(
            model=st.session_state.get('model', 'deepseek-chat'),
            messages=[
                {"role": "system", "content": "你是一个专业的软件需求分析师，能够准确识别文档中的软件需求。"},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3
        )
        
        result_text = response.choices[0].message.content
        
        # 解析结果
        import json
        json_match = re.search(r'\{.*\}', result_text, re.DOTALL)
        if json_match:
            result = json.loads(json_match.group())
            requirements = result.get("requirements", [])
            
            # 过滤空需求和过短需求
            filtered_reqs = [req.strip() for req in requirements 
                           if req.strip() and len(req.strip()) > MIN_PARAGRAPH_LENGTH]
            
            return filtered_reqs
        else:
            st.warning("AI需求识别返回格式不正确")
            return []
            
    except Exception as e:
        logger.error(f"AI需求识别失败 ({filename}): {e}")
        st.warning(f"AI需求识别失败: {str(e)}")
        return []